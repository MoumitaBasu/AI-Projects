def run_audio_to_text():
    import streamlit as st
    import whisper
    from streamlit_mic_recorder import mic_recorder, speech_to_text
    from googletrans import Translator, LANGUAGES
    from docx import Document
    from fpdf import FPDF
    import tempfile
    import os
    import re
    from google.cloud import speech
    import io

    #st.set_page_config(page_title="Whisper Translate", layout="centered")
    #st.title("🎤 Whisper Translator App")

    # ---------- Font Map and Script Detection (same as your provided code) ----------
    SCRIPT_FONT_MAP = {
        "Arabic": ("NotoSansArabic", "fonts/NotoSansArabic-Regular.ttf"),
        "Bengali": ("NotoSansBengali", "fonts/NotoSansBengali-Regular.ttf"),
        "Tamil": ("NotoSansTamil", "fonts/NotoSansTamil-Regular.ttf"),
        "Telugu": ("NotoSansTelugu", "fonts/NotoSansTelugu-Regular.ttf"),
        "Kannada": ("NotoSansKannada", "fonts/NotoSansKannada-Regular.ttf"),
        "Malayalam": ("NotoSansMalayalam", "fonts/NotoSansMalayalam-Regular.ttf"),
        "Thai": ("NotoSansThai", "fonts/NotoSansThai-Regular.ttf"),
        "Khmer": ("NotoSansKhmer", "fonts/NotoSansKhmer-Regular.ttf"),
        "Sinhala": ("NotoSansSinhala", "fonts/NotoSansSinhala-Regular.ttf"),
        "Georgian": ("NotoSansGeorgian", "fonts/NotoSansGeorgian-Regular.ttf"),
        "Armenian": ("NotoSansArmenian", "fonts/NotoSansArmenian-Regular.ttf"),
        "CJK-SC": ("NotoSansCJK", "fonts/NotoSansCJK-Regular.ttc", 0),
        "CJK-JP": ("NotoSansCJK", "fonts/NotoSansCJK-Regular.ttc", 1),
        "CJK-KR": ("NotoSansCJK", "fonts/NotoSansCJK-Regular.ttc", 2),
        "CJK-TC": ("NotoSansCJK", "fonts/NotoSansCJK-Regular.ttc", 3),
        "Default": ("NotoSans", "fonts/NotoSans-Regular.ttf")
    }

    SCRIPT_PATTERNS = {
        "Arabic": r"[\u0600-\u06FF]",
        "Bengali": r"[\u0980-\u09FF]",
        "Tamil": r"[\u0B80-\u0BFF]",
        "Telugu": r"[\u0C00-\u0C7F]",
        "Kannada": r"[\u0C80-\u0CFF]",
        "Malayalam": r"[\u0D00-\u0D7F]",
        "Thai": r"[\u0E00-\u0E7F]",
        "Khmer": r"[\u1780-\u17FF]",
        "Sinhala": r"[\u0D80-\u0DFF]",
        "Georgian": r"[\u10A0-\u10FF]",
        "Armenian": r"[\u0530-\u058F]",
        "CJK": r"[\u4E00-\u9FFF]",
    }

    def detect_script(text):
        if re.search(r"[繁體臺灣香港萬學愛藝]", text):
            return "CJK-TC"
        if re.search(r"[\u4E00-\u9FFF]", text): return "CJK-SC"
        if re.search(r"[\u3040-\u30FF]", text): return "CJK-JP"
        if re.search(r"[\uAC00-\uD7AF]", text): return "CJK-KR"
        for script, pattern in SCRIPT_PATTERNS.items():
            if re.search(pattern, text):
                return script
        return "Default"

    # ---------- Language options for speech recognition ----------
    language_options = {
        "Afrikaans": "af", "Amharic": "am", "Arabic": "ar", "Armenian": "hy", "Azerbaijani": "az",
        "Basque": "eu", "Bengali": "bn", "Bosnian": "bs", "Bulgarian": "bg", "Catalan": "ca",
        "Chinese (Mandarin)": "zh-CN", "Chinese (Cantonese)": "zh-HK", "Chinese (Taiwanese Mandarin)": "zh-TW",
        "Croatian": "hr", "Czech": "cs", "Danish": "da", "Dutch": "nl", "English (US)": "en-US",
        "English (UK)": "en-GB", "Estonian": "et", "Filipino": "fil", "Finnish": "fi", "French": "fr",
        "Galician": "gl", "Georgian": "ka", "German": "de", "Greek": "el", "Gujarati": "gu",
        "Hebrew": "he", "Hindi": "hi", "Hungarian": "hu", "Icelandic": "is", "Indonesian": "id",
        "Italian": "it", "Japanese": "ja", "Javanese": "jv", "Kannada": "kn", "Kazakh": "kk",
        "Khmer": "km", "Korean": "ko", "Lao": "lo", "Latvian": "lv", "Lithuanian": "lt",
        "Macedonian": "mk", "Malay": "ms", "Malayalam": "ml", "Marathi": "mr", "Mongolian": "mn",
        "Nepali": "ne", "Norwegian": "no", "Persian": "fa", "Polish": "pl",
        "Portuguese (Portugal)": "pt-PT", "Portuguese (Brazil)": "pt-BR", "Punjabi": "pa",
        "Romanian": "ro", "Russian": "ru", "Serbian": "sr", "Sinhala": "si", "Slovak": "sk",
        "Slovenian": "sl", "Spanish (Spain)": "es-ES", "Spanish (Mexico)": "es-MX",
        "Swahili": "sw", "Swedish": "sv", "Tamil": "ta", "Telugu": "te", "Thai": "th",
        "Turkish": "tr", "Ukrainian": "uk", "Urdu": "ur", "Vietnamese": "vi", "Zulu": "zu"
    }
    # Google Cloud Speech Client
    def transcribe_with_gemini(audio_file):
        client = speech.SpeechClient()
        audio = speech.RecognitionAudio(content=audio_file)
        
        config = speech.RecognitionConfig(
            encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
            sample_rate_hertz=16000,
            language_code="en-US",  # You can modify this based on language detection
        )

        response = client.recognize(config=config, audio=audio)
        
        # Extract the transcribed text
        transcription = ""
        for result in response.results:
            transcription += result.alternatives[0].transcript
        return transcription

    # ---------- Mode selection ----------
    mode = st.radio("Choose Input Mode", ["Upload Audio", "Record Audio"])

    # ---------- Upload Audio Mode ----------
    if mode == "Upload Audio":
        audio_file = st.file_uploader("Upload Audio", type=["wav", "mp3", "m4a"])
        if audio_file:
            st.audio(audio_file)

            model = whisper.load_model("base")
            if st.button("Transcribe Audio"):
                with tempfile.NamedTemporaryFile(delete=False) as tmp:
                    tmp.write(audio_file.read())
                    tmp_path = tmp.name

                result = model.transcribe(tmp_path)
                os.remove(tmp_path)
                transcription = result["text"]
                st.session_state.transcription_text = transcription
                st.markdown("### Transcribed Text")
                st.write(transcription)

    # ---------- Record Audio Mode ----------
    if mode == "Record Audio":
        selected_language_label = st.selectbox("Select language for speech recognition", list(language_options.keys()))
        selected_language_code = language_options[selected_language_label]

        st.write("🎙️ Record your voice:")

        # Record the audio
        audio_data = mic_recorder(record=True, language=selected_language_code)
        
        if audio_data:
            # Auto detect and transcribe with Gemini (Google Cloud Speech API)
            transcription = transcribe_with_gemini(audio_data)
            st.session_state.transcription_text = transcription
            st.markdown("### Transcribed Text")
            st.write(transcription)

    # ---------- Translation ----------
    if "transcription_text" in st.session_state and st.session_state.transcription_text:
        translator = Translator()
        st.markdown("---")
        st.subheader("🌍 Translate Text")

        lang_options = {"Keep Original": "original"}
        for code, name in LANGUAGES.items():
            lang_options[name.title()] = code
        lang_options = dict(sorted(lang_options.items()))

        target_lang = st.selectbox("Translate to", options=list(lang_options.keys()))
        if target_lang != "Keep Original":
            translation = translator.translate(st.session_state.transcription_text, dest=lang_options[target_lang])
            st.session_state.translated_text = translation.text
        else:
            st.session_state.translated_text = st.session_state.transcription_text

        st.markdown("### Translated Text")
        st.write(st.session_state.translated_text)

    # ---------- Download Options ----------
    if "translated_text" in st.session_state:
        st.markdown("---")
        st.subheader("⬇️ Download Options")

        selections = st.multiselect("Select content to download", ["Original Text", "Translated Text"], default=["Translated Text"])
        file_format = st.selectbox("File format", ["DOCX", "PDF"])

        def save_docx(title, content):
            doc = Document()
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(tmp.name)
            return tmp.name

        def save_pdf(title, content):
            pdf = FPDF()
            pdf.add_page()
            script = detect_script(content)
            font_name, font_path = SCRIPT_FONT_MAP.get(script, SCRIPT_FONT_MAP["Default"])[:2]
            pdf.add_font(font_name, "", font_path, uni=True)
            pdf.set_font(font_name, size=14)
            pdf.multi_cell(0, 10, f"{title}\n\n{content}")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            pdf.output(tmp.name)
            return tmp.name

        if st.button("Generate File(s)"):
            if "Original Text" in selections:
                content = st.session_state.transcription_text
                title = "Original Transcription"
                path = save_docx(title, content) if file_format == "DOCX" else save_pdf(title, content)
                st.download_button(f"Download {title}", data=open(path, "rb").read(), file_name=f"{title}.{file_format.lower()}")

            if "Translated Text" in selections:
                content = st.session_state.translated_text
                title = "Translated Text"
                path = save_docx(title, content) if file_format == "DOCX" else save_pdf(title, content)
                st.download_button(f"Download {title}", data=open(path, "rb").read(), file_name=f"{title}.{file_format.lower()}")
